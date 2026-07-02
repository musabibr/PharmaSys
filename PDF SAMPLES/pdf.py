import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_perfect_rtl_document():
    doc = docx.Document()

    # 1. ضبط إعدادات الصفحة والهوامش
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # دالة سحرية لفرض الـ RTL على مستوى الفقرة بالكامل والتحكم في اتجاه الرموز
    def apply_strict_rtl(paragraph):
        paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        pPr = paragraph._element.get_or_add_pPr()
        
        # فرض اتجاه النص العربي (Right-to-Left)
        bidi = OxmlElement('w:bidi')
        bidi.set(qn('w:val'), '1')
        pPr.append(bidi)

    # دالة لإضافة نصوص عربية محمية التنسيق
    def add_arabic_run(paragraph, text, bold=False, size_pt=11, color=None):
        run = paragraph.add_run(text)
        run.bold = bold
        if color:
            run.font.color.rgb = color
            
        rPr = run._element.get_or_add_rPr()
        
        # إعلام الوورد أن هذا النص عبارة عن Complex Script (عربي)
        cs = OxmlElement('w:cs')
        cs.set(qn('w:val'), '1')
        rPr.append(cs)
        
        rtl = OxmlElement('w:rtl')
        rtl.set(qn('w:val'), '1')
        rPr.append(rtl)
        
        # تعيين نوع الخط الافتراضي للعربي
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Arial')
        rFonts.set(qn('w:hAnsi'), 'Arial')
        rFonts.set(qn('w:cs'), 'Arial')
        rPr.append(rFonts)
        
        # ضبط الحجم
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(size_pt * 2)))
        rPr.append(sz)
        szCs = OxmlElement('w:szCs')
        szCs.set(qn('w:val'), str(int(size_pt * 2)))
        rPr.append(szCs)
        
        return run

    # دالة مخصصة لحماية الكلمات الإنجليزية والأقواس من الانعكاس داخل السطر العربي
    def add_english_run(paragraph, text, size_pt=10.5):
        run = paragraph.add_run(text)
        rPr = run._element.get_or_add_rPr()
        
        # إجبار الوورد على معاملة هذا الجزء كـ LTR نقي لمنع قفز الأقواس والفاصلة
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Consolas')
        rFonts.set(qn('w:hAnsi'), 'Consolas')
        rFonts.set(qn('w:cs'), 'Consolas')
        rPr.append(rFonts)
        
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), str(int(size_pt * 2)))
        rPr.append(sz)
        
        # إضافة تظليل رمادي خفيف لتبدو الملاحظات اللاتينية واضحة ومعزولة
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), 'EAEAEA')
        rPr.append(shd)
        
        return run

    def start_bullet_paragraph():
        p = doc.add_paragraph()
        apply_strict_rtl(p)
        p.paragraph_format.right_indent = Inches(0.25)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.25
        add_arabic_run(p, "•  ", bold=False)
        return p

    # --- كتابة محتوى المستند ---
    
    # العنوان الرئيسي
    p_title = doc.add_paragraph()
    apply_strict_rtl(p_title)
    p_title.paragraph_format.space_after = Pt(18)
    add_arabic_run(p_title, "قائمة المشاكل الموجودة في الملف", bold=True, size_pt=18, color=RGBColor(31, 78, 121))

    # القسم الأول
    p_h1 = doc.add_paragraph()
    apply_strict_rtl(p_h1)
    p_h1.paragraph_format.space_before = Pt(12)
    p_h1.paragraph_format.space_after = Pt(8)
    add_arabic_run(p_h1, "أولاً: مشاكل التنسيق والتنظيم (الشكل)", bold=True, size_pt=14, color=RGBColor(43, 108, 176))

    # النقاط الخاصة بالقسم الأول
    p = start_bullet_paragraph()
    add_arabic_run(p, "المسافات في عناوين الأقسام: ", bold=True)
    add_arabic_run(p, "أرقام الأقسام ملتصقة بالعناوين بدون مسافة، مثل «")
    add_english_run(p, "1.1Introduction")
    add_arabic_run(p, "» و«")
    add_english_run(p, "3.2Research Philosophy")
    add_arabic_run(p, "». الصحيح هو ترك مسافة لتصبح: «")
    add_english_run(p, "1.1 Introduction")
    add_arabic_run(p, "».")

    p = start_bullet_paragraph()
    add_arabic_run(p, "أخطاء في صيغة الاقتباسات داخل النص: ", bold=True)
    add_arabic_run(p, "الاقتباس «")
    add_english_run(p, "(Wardatun, , M, T, 2024)")
    add_arabic_run(p, "» فيه فاصلة مزدوجة وصيغة اسم غير سليمة، وكذلك «")
    add_english_run(p, "et Al.")
    add_arabic_run(p, "» يجب أن تُكتب بحرف صغير «")
    add_english_run(p, "et al.")
    add_arabic_run(p, "».")

    p = start_bullet_paragraph()
    add_arabic_run(p, "الفقرة الأولى في المقدمة: ", bold=True)
    add_arabic_run(p, "ملتصقة بالفقرة التي تليها بسبب علامات جدولة (")
    add_english_run(p, "Tabs")
    add_arabic_run(p, ") ومسافات عشوائية محشورة في المنتصف.")

    p = start_bullet_paragraph()
    add_arabic_run(p, "نقطة زائدة في تذييل الصفحات (")
    add_english_run(p, "Footer")
    add_arabic_run(p, "): ", bold=True)
    add_arabic_run(p, "تظهر نقطة «.» قبل رقم الصفحة في الصفحات الأولى (تبدو مثل «")
    add_english_run(p, ".1")
    add_arabic_run(p, "» و«")
    add_english_run(p, ".I")
    add_arabic_run(p, "»).")

    p = start_bullet_paragraph()
    add_arabic_run(p, "جدول الفرضيات (4.8): ", bold=True)
    add_arabic_run(p, "الأعمدة ضيّقة جداً فتنكسر الكلمات في منتصفها، مثل «")
    add_english_run(p, "signifi cant")
    add_arabic_run(p, "» و«")
    add_english_run(p, "Regres sion")
    add_arabic_run(p, "» و«")
    add_english_run(p, "Not Suppor ted")
    add_arabic_run(p, "»، وهذا يجعل الجدول يبدو غير احترافي.")

    p = start_bullet_paragraph()
    add_arabic_run(p, "قائمة المراجع (")
    add_english_run(p, "References")
    add_arabic_run(p, "): ", bold=True)
    add_arabic_run(p, "غير مرتّبة أبجدياً، ومقسّمة إلى كتلتين بتنسيق ومسافات بادئة مختلفة (1–10 ثم 11–19)، ويوجد مرجع مكرّر (")
    add_english_run(p, "Romney & Steinbart, 2020")
    add_arabic_run(p, " مذكور مرتين: رقم 10 ورقم 15)، والترقيم والخط غير موحّدين.")

    p = start_bullet_paragraph()
    add_arabic_run(p, "ترقيم مزدوج في الاستبيان: ", bold=True)
    add_arabic_run(p, "أول سؤال مكتوب بطريقة خاطئة: «")
    add_english_run(p, "1. 1. Select your organization's sector")
    add_arabic_run(p, "».")

    p = start_bullet_paragraph()
    add_arabic_run(p, "صفحة الغلاف: ", bold=True)
    add_arabic_run(p, "عنوان البحث غير بارز (بنفس حجم عناوين الجامعة)، ولا يوجد تاريخ/سنة التقديم.")

    # القسم الثاني
    p_h2 = doc.add_paragraph()
    apply_strict_rtl(p_h2)
    p_h2.paragraph_format.space_before = Pt(16)
    p_h2.paragraph_format.space_after = Pt(8)
    add_arabic_run(p_h2, "ثانياً: مشاكل المحتوى والتوثيق (الجوهر)", bold=True, size_pt=14, color=RGBColor(43, 108, 176))

    # ملاحظة منبثقة
    p_note = doc.add_paragraph()
    apply_strict_rtl(p_note)
    p_note.paragraph_format.right_indent = Inches(0.4)
    p_note.paragraph_format.space_after = Pt(10)
    add_arabic_run(p_note, "تنبيه: هذه أمور قد يلاحظها المشرف، لكن بعضها يحتاج قراراً منكم لأنه يخص مضمون البحث، ولن أغيّره من تلقاء نفسي.", size_pt=10, color=RGBColor(120, 120, 120))

    # نقاط القسم الثاني
    p = start_bullet_paragraph()
    add_arabic_run(p, "مصدر مذكور في النص وغير موجود في القائمة: ", bold=True)
    add_arabic_run(p, "مثل «")
    add_english_run(p, "(Hall, 2015)")
    add_arabic_run(p, "». كل مصدر يُذكر في النص يجب أن يكون في قائمة المراجع — هذا تحتاجون إضافته بأنفسكم لأنني لا أعرف بياناته الكاملة.")

    p = start_bullet_paragraph()
    add_arabic_run(p, "حجم العينة مقابل التحليل: ", bold=True)
    add_arabic_run(p, "عدد المشاركين 52 فقط، بينما خطة التحليل تذكر «")
    add_english_run(p, "Factor Analysis / PCA")
    add_arabic_run(p, "» وانحداراً متعدداً بستة متغيرات. هذا العدد صغير نسبياً لهذه التحليلات، وقد يطلب المشرف تبريراً أو تخفيف بعض الاستنتاجات.")

    p = start_bullet_paragraph()
    add_arabic_run(p, "تفسير النتائج: ", bold=True)
    add_arabic_run(p, "بُعد «")
    add_english_run(p, "Timeliness")
    add_arabic_run(p, "» فقط هو المعنوي في الانحدار، ومع ذلك اعتُبرت الفرضية الرئيسية «مدعومة». تفسيركم (الارتباط العالي بين الأبعاد) منطقي، لكن يجب صياغته بوضوح ودقة لأنه نقطة قد يناقشها المشرف بشدة.")

    doc.save("Perfect_RTL_Report.docx")

if __name__ == "__main__":
    create_perfect_rtl_document()